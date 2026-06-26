"""
document_loader.py — Load PDF, DOCX, and TXT files into LangChain Documents
"""

import os
import tempfile
from pathlib import Path
from typing import List

from langchain_core.documents import Document


def load_pdf(file_path: str) -> List[Document]:
    """Load a PDF file using pypdf (no deprecated loaders)."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    documents = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={"page": i + 1, "file_type": "pdf"},
                )
            )
    return documents


def load_docx(file_path: str) -> List[Document]:
    """Load a DOCX file using python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(page_content=full_text, metadata={"file_type": "docx"})]


def load_txt(file_path: str) -> List[Document]:
    """Load a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(page_content=text, metadata={"file_type": "txt"})]


def load_document(uploaded_file) -> List[Document]:
    """
    Master loader — detects file type and delegates to the right loader.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        List of LangChain Document objects
    """
    file_name = uploaded_file.name
    suffix = Path(file_name).suffix.lower()

    # Write to a temp file so loaders can read from disk
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    try:
        if suffix == ".pdf":
            docs = load_pdf(tmp_path)
        elif suffix == ".docx":
            docs = load_docx(tmp_path)
        elif suffix == ".txt":
            docs = load_txt(tmp_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    finally:
        os.unlink(tmp_path)  # Clean up temp file

    if not docs:
        raise ValueError(f"No text could be extracted from '{file_name}'.")

    return docs
